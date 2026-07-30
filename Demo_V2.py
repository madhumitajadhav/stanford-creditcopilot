import pickle
import json
import csv
import docx
import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt

from dataclasses import dataclass
from typing import Literal
from langchain import OpenAI
from langchain.callbacks import get_openai_callback
from langchain.chains import ConversationChain
from langchain.chains.conversation.memory import ConversationSummaryMemory
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression
from sklearn.tree import DecisionTreeClassifier
from sklearn.svm import SVC
from sklearn.metrics import mean_squared_error, mean_absolute_error
from sklearn.metrics import classification_report, confusion_matrix
from sklearn.naive_bayes import MultinomialNB
from init_ver_1_11 import *
from questions_v2 import *
from datetime import datetime
from pathlib import Path


st.title("Credit Co-Pilot 🤖")


tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs(["Discovery >", "Analysis >", "Proposal >", "Credit Memo >", "LLM Based Credit Decision","ML Based Credit Decision"])
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')


FILE_PATH = "LLM_Response/"


def write_into_doc(filename, text):
    if (Path(FILE_PATH+filename)).is_file():
        doc = docx.Document(FILE_PATH+filename)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        txt = '\n'.join(fullText)
        
        p = doc.add_paragraph()
        runtime = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        run = p.add_run(runtime)
        run.font.name = 'Arial'
        run.font.size = docx.shared.Pt(12)
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = docx.shared.Pt(12)
        
        # Save the document
        doc.save(FILE_PATH+filename)
        
    else:
        doc = docx.Document()
        # Add a paragraph to the document
        p = doc.add_paragraph()
        
        # Add some formatting to the paragraph
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_after = 0
         
        # Add a run and format it
        runtime = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        run = p.add_run(runtime)
        run.font.name = 'Arial'
        run.font.size = docx.shared.Pt(12)
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = docx.shared.Pt(12)
        
        # Save the document
        doc.save(FILE_PATH+filename)

with tab2:
    st.header("Discovery!")
    start_time = datetime.now()
    questions_ls = list(discovery_questions.keys())        
    question = [questions_ls[0], questions_ls[1], questions_ls[2]]
    selected_question = st.selectbox("Select The Question", question)

    txt = st.text_area("Selected Full Question :",discovery_questions[selected_question],height=200)
    
    if st.button('Summon The Discovery Co-Pilot'):
        with st.spinner(text="This may take a moment..."):
            load_dotenv()
            genai.configure(api_key=GOOGLE_API_KEY)
              
            new_db = init()
            
            responses = user_input(txt, new_db)
            st.write(responses["output_text"])
            st.session_state.discovery_flag = 1
            st.session_state.discovery_response = responses["output_text"]
            write_into_doc("Discovery.docx","QUESTION:"+ txt +"ANSWER :"+ st.session_state.discovery_response)
            
    end_time = datetime.now()
    st.write("time ----->",end_time-start_time)    
        
with tab3:
    st.header("Analysis")
    start_time = datetime.now()
    
    analysis_quest_ls = list(analysis_questions.keys())
    analysis_quest = [analysis_quest_ls[0], analysis_quest_ls[1],analysis_quest_ls[2],analysis_quest_ls[3]]
    selected_question = st.selectbox("Select The Question",analysis_quest)

    txt = st.text_area("Selected Full Question :",analysis_questions[selected_question],height=200)
    
    if st.button('Summon The Analysis Co-Pilot'):
        try:
            if st.session_state.discovery_flag==1:
                with st.spinner(text="This may take a moment..."):
                    load_dotenv()
                    genai.configure(api_key=GOOGLE_API_KEY)
                      
                    new_db = init()
                    
                    responses = user_input(txt, new_db)
                    st.session_state.analysis_flag=1
                    st.write(responses["output_text"])
                    st.session_state.analysis_response = responses["output_text"]
                    write_into_doc("Analysis.docx","QUESTION:"+ txt +"ANSWER :"+ st.session_state.analysis_response)
            else:
                st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete discovery phase!</h5>", unsafe_allow_html=True)
        except:
          st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete discovery phase!</h5>", unsafe_allow_html=True)
    end_time = datetime.now()
    st.write("time ----->",end_time-start_time)    
     

with tab4:
    st.header("Proposal")
    start_time = datetime.now()
    
    proposal_quest_ls = list(proposal_questions.keys())
    proposal_quest = [proposal_quest_ls[0], proposal_quest_ls[1]]
    selected_question = st.selectbox("Select The Question",proposal_quest)

    txt = st.text_area("Selected Full Question :",proposal_questions[selected_question],height=200)
    if st.button('Summon The Proposal Co-Pilot'):
        try:
            if st.session_state.discovery_flag==1:
                try:
                    if st.session_state.analysis_flag==1:
                        with st.spinner(text="This may take a moment..."):
                            load_dotenv()
                            genai.configure(api_key=GOOGLE_API_KEY)
                              
                            new_db = init()
                            txt = st.session_state.discovery_response + st.session_state.analysis_response + " use all the above responses from the discovery phase and analysis phase and provide " + txt
                    
                            responses = user_input(txt, new_db)
                            st.session_state.proposal_flag = 1
                            st.session_state.proposal_response = responses["output_text"]
                            write_into_doc("Proposal.docx","QUESTION:"+ txt +"ANSWER :"+ st.session_state.proposal_response)
                            st.write(responses["output_text"])
                    else:
                        st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete analysis phase!</h5>", unsafe_allow_html=True)
                except:
                    st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete analysis phase!</h5>", unsafe_allow_html=True)
            else:
                st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete discovery phase!</h5>", unsafe_allow_html=True)
        except:
          st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete discovery phase!</h5>", unsafe_allow_html=True)

    end_time = datetime.now()
    st.write("time ----->",end_time-start_time)    
     

with tab5:
    st.header("Credit Memo")
    start_time = datetime.now()
    
    try:
        if st.session_state.discovery_flag==1:
            try:
                if st.session_state.analysis_flag==1:
                    try:
                        if st.session_state.proposal_flag==1:
                            credit_memo_questions = ("""You are credit underwriter focused on High Net Worth client credit. Generate a credit memorandum for Barbie based on the following information. Be detailed, factual and verbose."""
                            + st.session_state.discovery_response + st.session_state.analysis_response + st.session_state.proposal_response)

                            txt = st.text_area("Prepopulated Question with Prior Stage Responses:",credit_memo_questions,height=300)
                            
                            if st.button('Generate Credit Memorandum'):
                                with st.spinner(text="This may take a moment..."):
                                    load_dotenv()
                                    genai.configure(api_key=GOOGLE_API_KEY)
                                      
                                    new_db = init()
                                    st.session_state.credit_memo_flag = 1
                                    responses = user_input(txt, new_db)
                                    st.session_state.credit_memo_response = responses["output_text"]
                                    write_into_doc("Credit_Memorandum.docx","QUESTION:"+ txt +"ANSWER :"+ st.session_state.credit_memo_response)
                                    st.write(responses["output_text"])
                            else:
                                print("in else ...")
                        else:
                            st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete Proposal phase!</h5>", unsafe_allow_html=True)
                    except:
                           st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete Proposal phase!</h5>", unsafe_allow_html=True) 
                else:
                    st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete Analysis phase!</h5>", unsafe_allow_html=True)
            except:
                st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete Analysis phase!</h5>", unsafe_allow_html=True)
        else:
            st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete discovery phase!</h5>", unsafe_allow_html=True)
    except:
        st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete discovery phase!</h5>", unsafe_allow_html=True)
    end_time = datetime.now()
    st.write("time ----->",end_time-start_time)    
   

with tab6:
    st.header("Credit Decision")
    start_time = datetime.now()
    
    try:
        if st.session_state.discovery_flag==1:
            try:
                if st.session_state.analysis_flag==1:
                    try:
                        if st.session_state.proposal_flag==1:
                            try:
                                st.write()
                                if st.session_state.credit_memo_flag==1:
                                    credit_review_questions = (
    """As a credit risk officer, your goal is 2 step:
    1: critique the credit memorandum which is provided here, review and affirm or challenge the proposed CRR rating, and identify any risks associated with extending credit to the borrower in the context of the Tailored Lending Underwiritng Guidelines. 
    2: Based on the response in 1, you have to make the final credit decision to approve, approve with conditions (need to specify what they are) or reject the loan application using chain of thought reasoning.
    Submitted Credit Memorandum Below:""" + st.session_state.credit_memo_response)
                                    
                                    txt = st.text_area("Prepopulated Question with Prior Stage Responses:",credit_review_questions,height=300)
                            
                                    if st.button('Get Credit Decision'):
                                        with st.spinner(text="This may take a moment..."):
                                            load_dotenv()
                                            genai.configure(api_key=GOOGLE_API_KEY)
                                              
                                            new_db = init()
                                            
                                            responses = user_input(txt, new_db)
                                            proposal_flag = 1
                                            st.session_state.credit_decision_response = responses["output_text"]
                                            write_into_doc("Credit_Decision.docx","QUESTION:"+ txt +"ANSWER :"+ st.session_state.credit_decision_response)
                                            st.write(responses["output_text"])
                                else:
                                    st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete Credit Memo phase!</h5>", unsafe_allow_html=True)
                            except:
                                st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete Credit Memo phase!</h5>", unsafe_allow_html=True)
                        else:
                            st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete Proposal phase!</h5>", unsafe_allow_html=True)
                    except:
                           st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete Proposal phase!</h5>", unsafe_allow_html=True) 
                else:
                    st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete Analysis phase!</h5>", unsafe_allow_html=True)
            except:
                st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete Analysis phase!</h5>", unsafe_allow_html=True)
        else:
            st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete discovery phase!</h5>", unsafe_allow_html=True)
    except:
        st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete discovery phase!</h5>", unsafe_allow_html=True)
    end_time = datetime.now()
    st.write("time ----->",end_time-start_time)    
     

with tab7:
    
    model_pkl_file = "loan_classifier_model.pkl"  
    with open(model_pkl_file, 'rb') as file:  
        model = pickle.load(file)
    try:
        if st.session_state.discovery_flag==1:
            json_questions = {'Generate a JSON': f"""Generate a dictionary formatted structure without a header for Barbie
                                based on the information provided here: {st.session_state.discovery_response} and
                                ensure that
                                each value name pair has exactly one value only,
                                for education only have the highest degree,
                                all numbers should be represented as numbers and not strings,
                                you should do the simple math to arrive at the value,
                                if information is not available you have to use 0,
                                if the value is true use 1
                                and it should have only the following name value pairs mentioned here:
                                Is_Gender_Female,
                                Is_Marital_Status_known,
                                Has_known_dependents,
                                Is_a_graduate,
                                Self_Employed,
                                Applicant_Income,
                                Coapplicant's_Income,
                                Amount_of_Loan,
                                Term_of_Loan_in_months,
                                Is_Credit_History_Available,
                                Is_an_Urban_Area,
                                FICO_Score,
                                Loan_Amount_divided_by_Total_Assets,
                                Liquid_Assets_in_cash,
                                Liabilities,
                                Net_Worth
                                """}
            st.header("Confusion Matrix: Evaluation of Tree-Based Model on Test Data")
            st.image("confusion_matrix.png", output_format="PNG", caption="Confusion Matrix on the Test Set")
            st.markdown("-----")
            st.header("ML Model Decision:")
            
            if st.button("Get ML Prediction"):
                load_dotenv()
                genai.configure(api_key=GOOGLE_API_KEY)
                new_db = init()
                                                
                json_response = user_input(json_questions['Generate a JSON'], new_db)
                st.write(json_response)
                dict = json.loads(json_response['output_text'])
                new_row = pd.DataFrame(dict, index=[0])
                
                headers =['Gender','Married','Dependents','Education','Self_Employed','Applicant_Income','Coapplicant_Income',
                          'Loan_Amount','Term','Credit_History','Area','FICO_Score','Loan_To_Value','Liquid_Assets','Total_Debt','Net_Worth']
                new_row.columns = headers
                st.write(new_row)
                if model.predict(new_row)>0.5:
                    st.write(f"Based on the Tree Based ML Model prediction Loan is Approved. ")
                else:
                    st.write(f"Based on the Tree Based ML Model prediction Loan is Rejected. ")
                
        else:
            st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete discovery phase!</h5>", unsafe_allow_html=True)
    except:
        st.markdown("<h5 style='text-align: justify; color: Red'>Need to complete discovery phase!</h5>", unsafe_allow_html=True)
   

# Footer
st.markdown("---")
st.text("© 2024 Credit Co-Pilot")